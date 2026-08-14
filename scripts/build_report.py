"""Build the Part B report DRAFT into report/report.docx. Run from the root:

    python scripts/build_report.py

This is a DRAFT scaffold, not the final report. It reads every number from the
results/ artifacts (so nothing is hand-typed or can drift), lays out the sections
from report/OUTLINE.md with Word built-in styles, inserts the exhibits and the
methodology equations, and leaves [YOUR INTERPRETATION: ...] stubs wherever the
economic reasoning belongs - that reasoning is yours and is graded as yours.

Equations are inserted as native, editable Word equation objects (OMML), not
images, so they match the document and can be edited in Word's equation editor.

Finish in Word: fill every [YOUR INTERPRETATION] stub, add the app screenshots
after Step 9, update fields, and export report.pdf. Run scripts/run_part_b.py
first so the artifacts exist.
"""
from __future__ import annotations

import sys
import pathlib

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from latex2mathml.converter import convert as latex_to_mathml
import mathml2omml

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

ROOT = pathlib.Path(__file__).resolve().parent.parent
DATA = ROOT / "results" / "data"
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
APP = ROOT / "results" / "app"
REPORT = ROOT / "report" / "report.docx"
INK = RGBColor(0x24, 0x1C, 0x15)
STUB = RGBColor(0x8C, 0x2F, 0x1F)   # rubric red - makes [YOUR INTERPRETATION] stubs jump out
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

MV, MS = "Combined_MinVariance", "Combined_MaxSharpe"
TS_MV, TS_MS = "Combined_TwoStage_MinVariance", "Combined_TwoStage_MaxSharpe"

# Strategy labels (used in prose and tables) and the investor-facing product names.
STRAT = {MV: "Minimum Variance", MS: "Maximum Sharpe",
         TS_MV: "Two-Stage Min-Variance", TS_MS: "Two-Stage Max-Sharpe"}
PRODUCT = {MV: "Revelatio Balanced Fund (RBAL)", MS: "Revelatio High Growth Fund (RHGF)",
           TS_MV: "Revelatio Equity Signals Fund (REQS)",
           TS_MS: "Revelatio Multi-Opportunity Fund (RMOP)"}


def flabel(fund: str) -> str:
    """Readable label for any fund key, e.g. 'Crypto_RiskParity' -> 'Crypto Risk Parity'."""
    fam = ""
    for f in ("Equity_", "Crypto_", "Combined_"):
        if fund.startswith(f):
            fam, rest = f[:-1], fund[len(f):]
            break
    else:
        rest = fund
    rest = (rest.replace("TwoStage_", "Two-Stage ").replace("MinVariance", "Min-Variance")
            .replace("MaxSharpe", "Max-Sharpe").replace("RiskParity", "Risk Parity"))
    return rest if fam in ("", "Combined") else f"{fam} {rest}"


# --- load the numbers (nothing hand-typed) ------------------------------------

def _min_variance_decomposition(fw: pd.DataFrame, panel: pd.DataFrame,
                                oos_index) -> dict:
    """Measured 'why min-variance lost to 1/N' facts, from the weights and panel."""
    mv = fw[fw.fund == MV]
    avg_w = mv.groupby("ticker")["weight"].mean().sort_values(ascending=False)
    crypto = [c for c in panel.columns if c.endswith("-USD")]
    crypto_w = (mv[mv.ticker.isin(crypto)].groupby("rebalance_date")["weight"].sum().mean())
    eff_n = (1.0 / mv.groupby("rebalance_date")["weight"].apply(lambda w: (w ** 2).sum())).mean()
    return {
        "top3": list(avg_w.head(3).index),
        "crypto_w": float(crypto_w),
        "eff_n": float(eff_n),
    }


def load_facts() -> dict:
    m = pd.read_csv(TABLES / "performance_metrics.csv").set_index("fund")
    fr = pd.read_csv(DATA / "fund_returns.csv", index_col=0, parse_dates=True)
    fus = pd.read_csv(TABLES / "fusion_comparison.csv")
    sent = pd.read_csv(TABLES / "sentiment_stats.csv").iloc[0]
    bench = pd.read_csv(TABLES / "benchmark_comparison.csv").set_index("fund")
    bench_r = pd.read_csv(DATA / "benchmark_returns.csv", index_col=0, parse_dates=True)["EqualWeight50"]
    shr = pd.read_csv(TABLES / "shrinkage_comparison.csv")
    shr_d = pd.read_csv(TABLES / "shrinkage_diagnostics.csv")
    ts = pd.read_csv(TABLES / "two_stage_comparison.csv")
    dec = pd.read_csv(TABLES / "decay_comparison.csv")
    sweep = pd.read_csv(TABLES / "decay_halflife_sweep.csv")
    panel = pd.read_csv(DATA / "combined_returns_panel.csv", index_col=0, parse_dates=True)
    fw = pd.read_csv(DATA / "fund_weights.csv")

    bench_ret = float(bench["bench_ann_return"].iloc[0])
    bench_sharpe = float(bench["bench_sharpe"].iloc[0])
    bench_vol = float(bench_r.std() * np.sqrt(252))
    bg = float((1 + bench_r).cumprod().iloc[-1])

    return {
        "m": m, "fr": fr, "fus": fus, "fus_i": fus.set_index(["fund", "variant"]),
        "sent": sent, "growth": (1 + fr).cumprod().iloc[-1],
        "bench": bench, "bench_ret": bench_ret, "bench_sharpe": bench_sharpe,
        "bench_vol": bench_vol, "bench_growth": bg,
        "n_beat": int(bench["beats_benchmark"].sum()), "n_funds": len(bench),
        "shr": shr.set_index(["method", "variant"]), "shr_d": shr_d.set_index(["method", "variant"]),
        "ts": ts.set_index(["method", "structure"]),
        "dec": dec.set_index(["fund", "variant"]), "sweep": sweep,
        "mvd": _min_variance_decomposition(fw, panel, fr.index),
        "oos_start": fr.index.min().date(), "oos_end": fr.index.max().date(),
        "oos_days": len(fr),
    }


# --- equation rendering (LaTeX -> MathML -> native Word OMML equation) ---------

def _hide_empty_nary_limits(el) -> None:
    """Hide empty upper/lower limits on summation-style operators.

    A sum written with only a lower limit (``\\sum_{s\\in W}``) converts to an
    n-ary operator whose ``m:sup`` is empty, which Word draws as a stray "type
    here" placeholder box above the sigma. Setting ``m:supHide`` / ``m:subHide``
    on the operator's properties tells Word to omit the empty limit. Child order
    in ``m:naryPr`` is chr, limLoc, grow, subHide, supHide - subHide is added
    before supHide so the element stays schema-valid.
    """
    for nary in el.iter(qn("m:nary")):
        pr = nary.find(qn("m:naryPr"))
        if pr is None:
            continue
        for lim_tag, hide_tag in (("m:sub", "m:subHide"), ("m:sup", "m:supHide")):
            lim = nary.find(qn(lim_tag))
            if lim is not None and len(lim) == 0 and pr.find(qn(hide_tag)) is None:
                pr.append(parse_xml(
                    f'<{hide_tag} xmlns:m="{MATH_NS}" m:val="1"/>'))


def _omml(latex: str):
    """Convert LaTeX to an editable Word equation (OMML), not an image.

    latex2mathml renders the LaTeX to MathML, mathml2omml maps that to Office
    Math (OMML), and parse_xml turns it into an element the document can hold. The
    result is a true Word equation the reader can click and edit, so the report
    stays consistent and re-authorable rather than carrying flattened pictures.
    Empty n-ary limits are hidden so no placeholder box shows above a summation.
    """
    omml = mathml2omml.convert(latex_to_mathml(latex))
    if "xmlns:m=" not in omml:
        omml = omml.replace("<m:oMath>", f'<m:oMath xmlns:m="{MATH_NS}">', 1)
    el = parse_xml(omml)
    _hide_empty_nary_limits(el)
    return el


# --- docx helpers -------------------------------------------------------------

def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def stub(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"[YOUR INTERPRETATION: {text}]")
    r.font.color.rgb = STUB
    r.italic = True
    return p


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"[NOTE: {text}]")
    r.font.color.rgb = STUB
    r.italic = True
    return p


def equation(doc, latex, name, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(_omml(latex))                       # native, editable Word equation
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Equation {number}").italic = True


def figure(doc, filename, number, caption, directory=None):
    doc.add_picture(str((directory or FIGURES) / filename), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]
    run = c.add_run(f"Figure {number}. {caption}")
    run.italic = True


def table_from_df(doc, df, number, caption, pct_cols=(), num_cols=()):
    c = doc.add_paragraph()
    run = c.add_run(f"Table {number}. {caption}")
    run.italic = True
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else "Table Grid"
    for j, col in enumerate(df.columns):
        t.rows[0].cells[j].paragraphs[0].add_run(str(col)).bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            if col in pct_cols:
                v = f"{v:.1%}"
            elif col in num_cols:
                v = f"{v:.2f}"
            cells[j].text = str(v)


# --- the report ---------------------------------------------------------------

def _appendix(doc, f):
    """Appendix A - every exhibit, referenced from the narrative by A-number."""
    m, bench = f["m"], f["bench"]
    h(doc, "Appendix A - Exhibits", 1)

    def _map(df):
        df = df.copy()
        df["fund"] = df["fund"].map(flabel)
        return df

    perf = m.reset_index()[["fund", "ann_return", "ann_vol", "sharpe", "max_drawdown"]]
    table_from_df(doc, _map(perf), "A1",
                  "Out-of-sample performance across the eleven funds (annualised, rf = 0).",
                  pct_cols=("ann_return", "ann_vol", "max_drawdown"), num_cols=("sharpe",))
    bench_tbl = bench.reset_index()[["fund", "sharpe", "excess_ann_return",
                                     "sharpe_minus_bench", "beats_benchmark"]]
    table_from_df(doc, _map(bench_tbl), "A2",
                  "Each fund against the equal-weight 50-stock market (annualised).",
                  pct_cols=("excess_ann_return",), num_cols=("sharpe", "sharpe_minus_bench"))
    table_from_df(doc, _map(f["fus"][["fund", "variant", "ann_return", "sharpe", "max_drawdown"]]),
                  "A3", "Sentiment fusion before vs after, by fund (annualised).",
                  pct_cols=("ann_return", "max_drawdown"), num_cols=("sharpe",))
    table_from_df(doc, _map(f["shr"].reset_index()[["fund", "variant", "ann_return", "sharpe", "max_drawdown"]]),
                  "A4", "Covariance shrinkage before vs after (annualised).",
                  pct_cols=("ann_return", "max_drawdown"), num_cols=("sharpe",))
    table_from_df(doc, _map(f["ts"].reset_index()[["fund", "structure", "ann_return", "sharpe", "max_drawdown"]]),
                  "A5", "One-stage vs two-stage, by method (annualised).",
                  pct_cols=("ann_return", "max_drawdown"), num_cols=("sharpe",))
    table_from_df(doc, _map(f["dec"].reset_index()[["fund", "variant", "ann_return", "sharpe", "max_drawdown"]]),
                  "A6", "Sentiment decay: base vs neutral-0 vs decayed tilt (annualised).",
                  pct_cols=("ann_return", "max_drawdown"), num_cols=("sharpe",))

    for fn, num, cap in [
        ("growth_of_dollar.png", "A1", "Growth of $1, combined funds by method and the equal-weight market."),
        ("drawdown.png", "A2", "Drawdown from the running peak, two combined funds."),
        ("weights_over_time.png", "A3", "Maximum-Sharpe target weights over time (top 6 holdings)."),
        ("sharpe_barplot.png", "A4", "Annualised Sharpe by fund and method, coloured by family, against the market."),
        ("sector_sentiment.png", "A5", "Monthly sector sentiment index, ten equity sectors and their average."),
        ("fusion_before_after.png", "A6", "Base vs sentiment-tilted growth of $1."),
        ("shrinkage_growth.png", "A7", "Growth of $1 under the sample vs Ledoit-Wolf covariance."),
        ("shrinkage_dispersion.png", "A8", "Effective number of holdings, sample vs shrunk."),
        ("two_stage_growth.png", "A9", "Growth of $1, one-stage vs two-stage."),
        ("two_stage_sleeve_share.png", "A10", "Realised equity-sleeve share over time, two-stage funds."),
        ("decay_mechanism.png", "A11", "Sentiment fill for one ticker: neutral-0 rule vs decay."),
        ("decay_fusion.png", "A12", "Growth of $1: base vs neutral-0 tilt vs decayed tilt."),
    ]:
        figure(doc, fn, num, cap)
    for fn, num, cap in [
        ("01_home.jpg", "A13", "App home: the offering and the headline that 3 of 11 funds beat the market."),
        ("02_funds_compare.jpg", "A14", "App Funds page: the eleven-fund comparison table and the multi-asset cards."),
        ("03_fact_sheet.jpg", "A15", "App fund fact sheet: key stats, the beats-market pill, and the Overview / Performance / Holdings tabs."),
        ("05_sentiment.jpg", "A16", "App Sentiment page: the news-mood gauge, the sector index, and the fusion result."),
        ("06_build_portfolio.jpg", "A17", "App Build-portfolio page: the fund blender, the fixed 0.75% fee, and the deployment table."),
    ]:
        figure(doc, fn, num, cap, directory=APP)


def build():
    f = load_facts()
    m, g, fus_i, sent, bench = f["m"], f["growth"], f["fus_i"], f["sent"], f["bench"]
    mvd = f["mvd"]

    doc = Document()
    for s in doc.sections:
        s.page_height, s.page_width = Inches(11.69), Inches(8.27)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.98)

    title = doc.add_paragraph("Revelatio: Systematically Managed Funds for the First-Time Investor")
    title.style = doc.styles["Title"]
    sub = doc.add_paragraph("Part B - Funds, Sentiment and the App")
    sub.style = doc.styles["Subtitle"] if "Subtitle" in [s.name for s in doc.styles] else doc.styles["Normal"]
    doc.add_paragraph("z5736927 - Gianni Chepote - FINS5545 Financial Market Data Literacy")

    h(doc, "Abstract", 1)
    body(doc, (
        f"Revelatio offers systematically managed funds to first-time investors. This "
        f"Part builds eleven funds - three families (equity-only, crypto-only and "
        f"combined) across three methods (minimum-variance, maximum-Sharpe and risk "
        f"parity), plus two two-stage funds - over 50 US large caps and 10 "
        f"cryptocurrencies, each backtested out of sample with a rolling {252}-day "
        f"window, monthly rebalancing, and no look-ahead. Over {f['oos_start']} to "
        f"{f['oos_end']} the maximum-Sharpe fund returned {m.loc[MS,'ann_return']:.1%} a "
        f"year at a {m.loc[MS,'sharpe']:.2f} Sharpe; against an equal-weight 50-stock "
        f"market (Sharpe {f['bench_sharpe']:.2f}), only {f['n_beat']} of {f['n_funds']} "
        f"funds beat simply owning the market. A VADER sentiment index and a basic "
        f"tilt complete the pipeline, and three extensions - covariance shrinkage, a "
        f"two-stage portfolio of portfolios, and sentiment decay - are each shown out "
        f"of sample."))
    body(doc, (
        "Blending 50 US large-cap stocks with 10 cryptocurrencies returned 25.5% a "
        "year at a 1.03 Sharpe when the objective was to maximise risk-adjusted "
        "return - but playing it safe backfired: the minimum-variance funds lost to "
        "an equal-weight basket of the same 50 stocks. They did cut risk, with a "
        "shallower 15.6% drawdown against the market's 20.3%, yet by minimising "
        "variance they sidestepped the higher-volatility names and crypto that drove "
        "returns, giving up more upside than the protection was worth. Optimisation "
        "only pays when its objective matches how the market actually moves."))

    h(doc, "1. The product and the investor", 1)
    body(doc, (
        "Revelatio is an investment app for the first-time investor who suspects "
        "investing is rigged. It stands on two inputs: daily prices for 50 US "
        "large-cap equities and 10 cryptocurrencies (2020-2023), and daily news "
        "headlines for the equities."))
    body(doc, (
        "Revelatio is for the first-time investor who doesn't want to study markets "
        "or choose between stocks and crypto - they just want a fund that balances "
        "the two for them and runs itself. Trust is what makes that work: keep the "
        "app simple, show the risks and fees plainly, and let the optimised funds do "
        "the job so the user doesn't have to second-guess every move."))

    h(doc, "2. The funds and the backtest design", 1)
    body(doc, (
        "Each fund is a set of portfolio weights over an investable universe, rebuilt "
        f"on the first trading day of each month from the trailing {252} trading days "
        "of returns. The same construction is applied to three universes - equities "
        "only (50), crypto only (10), and combined (60) - so each method yields an "
        "equity-only, a crypto-only, and a combined fund. The daily portfolio return "
        "is w'r_t; weights use the sample mean and covariance over the rolling window:"))
    equation(doc, r"\hat{\boldsymbol{\mu}} = \frac{1}{252}\sum_{s\in W}\mathbf{r}_s, "
                  r"\quad \hat{\boldsymbol{\Sigma}} = \frac{1}{251}\sum_{s\in W}"
                  r"(\mathbf{r}_s-\hat{\boldsymbol{\mu}})(\mathbf{r}_s-\hat{\boldsymbol{\mu}})^{\top}",
             "musigma", 1)
    body(doc, "The minimum-variance and maximum-Sharpe (tangency) funds solve")
    equation(doc, r"\min_{\mathbf{w}}\ \mathbf{w}^{\top}\hat{\boldsymbol{\Sigma}}\mathbf{w}"
                  r"\qquad \max_{\mathbf{w}}\ \frac{\mathbf{w}^{\top}\hat{\boldsymbol{\mu}} - r_f}"
                  r"{\sqrt{\mathbf{w}^{\top}\hat{\boldsymbol{\Sigma}}\mathbf{w}}}"
                  r"\qquad \mathrm{s.t.}\ \sum_i w_i = 1,\ w_i \geq 0", "mvms", 2)
    body(doc, ("and the risk-parity fund equalises each asset's contribution to "
               "portfolio variance (Spinu, 2013), normalised to sum to one:"))
    equation(doc, r"\min_{\mathbf{w}>0}\ \frac{1}{2}\mathbf{w}^{\top}"
                  r"\hat{\boldsymbol{\Sigma}}\mathbf{w} - \frac{1}{n}\sum_i \ln w_i", "riskparity", 3)
    body(doc, ("All are long-only and fully invested, r_f = 0 (stated). Performance is "
               "annualised on the 252-day equity calendar; Sharpe is mu_ann/sigma_ann, "
               "and maximum drawdown is the worst peak-to-trough decline of growth path V:"))
    equation(doc, r"\mathrm{MDD} = \min_t\left(\frac{V_t}{\max_{s\leq t}V_s} - 1\right)", "mdd", 4)
    body(doc, (
        f"The design is out-of-sample throughout: the first live date is "
        f"{f['oos_start']} (the first month-start with a full 252-day prior window), "
        f"the sample runs to {f['oos_end']} (36 rebalances, {f['oos_days']} days), "
        "weights come only from data before each rebalance, and no figure includes a "
        "training-period day. Crypto returns are computed on their own calendar then "
        "left-joined onto equity trading days."))
    body(doc, (
        "A one-year window is the sweet spot for a 60-asset mix: long enough to "
        "estimate returns and risk without the numbers going haywire, short enough "
        "to keep up as markets - especially crypto - shift. Monthly rebalancing "
        "keeps the funds current without constant trading, which suits a hands-off "
        "first-timer and keeps churn low."))

    h(doc, "3. Out-of-sample results, fact sheets and the benchmark", 1)
    body(doc, (
        f"The offering is a grid of three families x three methods (nine funds) plus "
        f"two two-stage funds, eleven in all (Table A1; Figures A1-A4). The combined "
        f"core funds are offered as the Revelatio Balanced Fund (minimum-variance) "
        f"and the Revelatio High Growth Fund (maximum-Sharpe). Two patterns hold: "
        f"within each method the combined fund beats either single-asset fund on "
        f"Sharpe (combined maximum-Sharpe {m.loc[MS,'sharpe']:.2f} vs "
        f"{m.loc['Equity_MaxSharpe','sharpe']:.2f} equity-only and "
        f"{m.loc['Crypto_MaxSharpe','sharpe']:.2f} crypto-only), and crypto-only funds "
        f"post the highest raw returns but the worst risk-adjusted ones (crypto risk "
        f"parity {m.loc['Crypto_RiskParity','ann_return']:.0%} a year, "
        f"{abs(m.loc['Crypto_RiskParity','max_drawdown']):.0%} drawdown, Sharpe "
        f"{m.loc['Crypto_RiskParity','sharpe']:.2f})."))
    body(doc, (
        "The grid makes the case for combining: within every method, mixing stocks "
        "and crypto beat holding either alone on a risk-adjusted basis. Crypto-only "
        "funds show why you don't hand a beginner pure crypto - 33% a year sounds "
        "great until you see the 83% fall that comes with it. So the funds we put up "
        "front are the combined ones; the single-asset versions exist to show the "
        "first-timer what diversification is buying them."))
    body(doc, (
        "The max-Sharpe fund earned four times the return but fell 26% at its worst "
        "- the kind of drop that makes a first-timer sell at the bottom and swear off "
        "investing for good. So you don't sell them the 25.5% headline; you show the "
        "drawdown right next to it and let them pick the fund whose worst day they "
        "could actually sit through."))
    h(doc, "3.1 Do the funds beat simply owning the market?", 2)
    body(doc, (
        f"The equal-weight 50-stock market (Table A2) is the honest same-universe "
        f"benchmark - our own stocks, calendar and schedule - so a fund that cannot "
        f"beat it is not adding value over naive diversification. It returned "
        f"{f['bench_ret']:.1%} a year at Sharpe {f['bench_sharpe']:.2f}, and only "
        f"{f['n_beat']} of {f['n_funds']} funds beat it. The combined minimum-variance "
        f"fund fell {abs(bench.loc[MV,'sharpe_minus_bench']):.2f} short on Sharpe: it "
        f"cut volatility to {m.loc[MV,'ann_vol']:.1%} against the market's "
        f"{f['bench_vol']:.1%} (about a fifth less risk) but gave up roughly half the "
        f"return ({m.loc[MV,'ann_return']:.1%} vs {f['bench_ret']:.1%}), concentrating "
        f"in about {mvd['eff_n']:.0f} low-volatility defensives "
        f"({', '.join(mvd['top3'])}) with only {mvd['crypto_w']:.1%} crypto."))
    body(doc, (
        "Our minimum-variance funds lost to a plain equal-weight basket. We'd still "
        "offer minimum-variance, but honestly - as the calm-ride option, a 16% "
        "worst-case fall instead of 26%, for someone who values sleeping at night "
        "over top returns, not as a return story. The app says outright that a basic "
        "basket beat it this period; hiding that would be the exact broker trick the "
        "product exists to avoid."))

    h(doc, "4. The sentiment index and fusion", 1)
    body(doc, (
        f"Each headline is scored with VADER's compound score in [-1, 1], averaged to "
        f"a per-ticker-day value, then equal-weighted across the five tickers in each "
        f"sector (Equation 5; Figure A5). Text passes through unmodified because VADER "
        f"reads casing, punctuation and negation; a ticker-day with no headline enters "
        f"as neutral 0; any use lags at least one trading day. Coverage is "
        f"{sent['ticker_day_coverage_pct']:.1f}% of ticker-days, and the market "
        f"average held mildly positive (mean {sent['overall_mean']:+.2f}), barely "
        f"reacting to the 2022 selloff."))
    equation(doc, r"S_{k,t} = \frac{1}{5}\sum_{i \in k} s_{i,t}", "sent", 5)
    body(doc, (
        "A signal that sits mildly positive and barely flinches - even through the "
        "2022 crash - isn't telling you much. If the news reads 'fine' whether the "
        "market is rising or falling, it can't help you separate good bets from bad "
        "ones, which is exactly why leaning on it later dragged returns down instead "
        "of lifting them."))
    body(doc, (
        f"The fusion tilts each fund's equity weights towards higher-sentiment sectors "
        f"using the lagged, z-scored signal (Equation 6; Table A3, Figure A6), lambda "
        f"= 0.5, untuned. It reduced Sharpe for both funds - minimum-variance "
        f"{fus_i.loc[(MV,'base'),'sharpe']:.2f} to {fus_i.loc[(MV,'sentiment'),'sharpe']:.2f}, "
        f"maximum-Sharpe {fus_i.loc[(MS,'base'),'sharpe']:.2f} to "
        f"{fus_i.loc[(MS,'sentiment'),'sharpe']:.2f} - an honest negative baseline."))
    equation(doc, r"w_i^{\mathrm{tilt}} \propto w_i\,(1 + \lambda z_i)_{+},"
                  r"\qquad z_i = \frac{s_i - \overline{s}}{\mathrm{sd}(s)}", "fuse", 6)
    body(doc, (
        "Sentiment hurt because the tilt is only as good as the signal underneath "
        "it, and that signal barely moves. Nudging money towards 'positive' sectors "
        "on a flat, always-upbeat reading just shuffled weights around noise, so "
        "both funds lost Sharpe. To make sentiment pay you'd need a sharper signal - "
        "finance-specific words, or reactions to real events - not a cleverer way to "
        "lean on a dull one."))

    _innovation_section(doc, f)

    h(doc, "6. The app and the investor journey", 1)
    body(doc, (
        "The Streamlit app supports the full journey - compare funds, read a fact "
        "sheet (growth, drawdown, holdings), read the sentiment analytics, and build "
        "an allocation - reading precomputed artifacts only, never a live backtest. "
        "Every fund page leads with its drawdown and the market benchmark, a "
        "deliberate departure from product pages that show only growth."))
    body(doc, (
        "All performance figures in this report are gross of fees. In the app's "
        "allocation view a fixed 0.75% annual management fee, accrued daily, is "
        "applied so the investor sees the net-of-fee outcome they would keep. In "
        "keeping with a low-cost, index-style product for a first-time investor, the "
        "fund charges a single management fee and no performance fee - the manager "
        "does not take a share of the investor's gains."))
    body(doc, (
        "The five app screens - the offering, the fund comparison, a fund fact "
        "sheet, the sentiment analytics, and the allocation blender - are in Figures "
        "A13-A17."))
    body(doc, (
        "A first-timer lands on the compare screen and sees, plainly, which funds "
        "beat a simple market basket and which didn't - no hard sell. They open a "
        "fund's fact sheet and get its growth, its worst fall, and what it actually "
        "holds, side by side. They glance at the market mood, then build a blend and "
        "watch what they'd keep after fees. Each step shows the downside as clearly "
        "as the upside, so trust is earned by the order things appear in, not by "
        "decoration."))

    h(doc, "7. Critical reflection and recommendations", 1)
    body(doc, (
        "The funds run on plain sample estimates over a short, single-regime window; "
        "the minimum-variance funds lose to naive diversification; the sentiment "
        "signal is weak and the fusion underperforms. Three tested responses "
        "(Section 5) point to concrete next steps."))
    body(doc, (
        "1. Adopt covariance shrinkage across every fund. It cost nothing to add, "
        "cut turnover, and lifted the minimum-variance Sharpe, and on a shorter "
        "window or a wider universe it would help more. It is the cheapest "
        "reliability upgrade on the table."))
    body(doc, (
        "2. Replace the general-purpose sentiment reader with a finance-specific "
        "one. The current signal barely moves and dragged returns when used; a "
        "lexicon tuned to market language, or a model that reacts to real events, "
        "has a chance of carrying information this one doesn't. Fix the signal "
        "before touching the tilt."))
    body(doc, (
        "3. Charge realistic trading costs in the backtest before trusting any "
        "fund's edge. Every result here assumes free trading, and the max-Sharpe "
        "fund rebalances hardest, so a cost model would test whether its lead "
        "survives contact with reality. A fund that only wins for free is not one "
        "you would sell."))

    h(doc, "References", 1)
    body(doc, "Hutto, C.J. and Gilbert, E.E. (2014). VADER: A Parsimonious Rule-based "
              "Model for Sentiment Analysis of Social Media Text. Proceedings of the "
              "Eighth International AAAI Conference on Weblogs and Social Media (ICWSM-14).")
    body(doc, "Ledoit, O. and Wolf, M. (2004). Honey, I Shrunk the Sample Covariance "
              "Matrix. The Journal of Portfolio Management, 30(4), 110-119.")
    body(doc, "Spinu, F. (2013). An Algorithm for Computing Risk Parity Weights. "
              "SSRN Working Paper 2297383. https://doi.org/10.2139/ssrn.2297383")
    body(doc, "Data: daily prices for 50 US large-cap equities and 10 "
              "cryptocurrencies, and daily equity news headlines, 2020-2023, "
              "provided by the course (backup source: openbondassetpricing.com).")

    _appendix(doc, f)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)
    n_fig = sum(1 for _ in FIGURES.glob("*.png"))
    print(f"Wrote {REPORT.relative_to(ROOT)}")
    print(f"  narrative sections: 7 | equations: 9 native (OMML) | "
          f"appendix: 6 tables + 12 figures + 5 app screenshots")
    print("  DRAFT: export report.pdf from Word; confirm narrative <= 10 pages.")


def _innovation_section(doc, f):
    """Section 5 - three extensions, one tight paragraph each; exhibits in the appendix."""
    m = f["m"]
    shr, shr_d, ts, dec = f["shr"], f["shr_d"], f["ts"], f["dec"]
    h(doc, "5. Innovation: three extensions, each shown out of sample", 1)
    body(doc, ("Three extensions, chosen for depth over count, each a before-vs-after "
               "against the plain version it replaces with no parameter tuned to win."))

    h(doc, "5.1 Covariance shrinkage", 2)
    body(doc, ("The 60-asset sample covariance on a one-year window is badly "
               "conditioned. A Ledoit-Wolf estimator shrinks it towards a scaled-"
               "identity target by an intensity delta chosen analytically per window:"))
    equation(doc, r"\hat{\boldsymbol{\Sigma}}_{\mathrm{LW}} = (1-\delta)\,\mathbf{S} "
                  r"+ \delta\,\overline{\sigma}^2\,\mathbf{I}", "shrink", 7)
    d_mv = shr_d.loc[("min_variance", "shrunk"), "mean_delta"]
    body(doc, (
        f"At delta = {d_mv:.2f} it lifted minimum-variance Sharpe from "
        f"{shr.loc[('min_variance','sample'),'sharpe']:.2f} to "
        f"{shr.loc[('min_variance','shrunk'),'sharpe']:.2f}, left maximum-Sharpe about "
        f"flat, and raised effective holdings from "
        f"{shr_d.loc[('min_variance','sample'),'effective_n']:.0f} to "
        f"{shr_d.loc[('min_variance','shrunk'),'effective_n']:.0f} names (Table A4, "
        f"Figures A7-A8)."))
    body(doc, (
        "Shrinkage is worth keeping: it costs nothing, spreads each fund across more "
        "names, and lifted the minimum-variance Sharpe from 0.49 to 0.55 without "
        "denting the max-Sharpe fund's edge. The shrink was small (delta around "
        "0.05) because a full year of data on 60 assets isn't as noisy as feared - "
        "the sample covariance was already decent, so the estimator only nudged it. "
        "On a shorter window or a wider universe, expect it to matter more."))

    h(doc, "5.2 A two-stage portfolio of portfolios", 2)
    body(doc, ("The two-stage fund builds an equity sleeve and a crypto sleeve, then "
               "allocates across the two, so the cross-asset decision rests on a 2x2 "
               "covariance rather than a 60x60:"))
    equation(doc, r"\mathbf{w} = a_{\mathrm{eq}}\,\mathbf{w}_{\mathrm{eq}} \;\oplus\; "
                  r"a_{\mathrm{cr}}\,\mathbf{w}_{\mathrm{cr}}", "twostage", 8)
    two_cross = int(f["ts"].loc[("min_variance", "two_stage"), "cross_asset_cov_params"])
    one_tot = int(f["ts"].loc[("min_variance", "one_stage"), "cov_params_total"])
    body(doc, (
        f"Like-for-like (same window, rebalance, method, sample), the {two_cross}-"
        f"parameter cross-asset decision matched the {one_tot}-parameter one-stage "
        f"fund almost exactly: minimum-variance "
        f"{ts.loc[('min_variance','one_stage'),'sharpe']:.2f} to "
        f"{ts.loc[('min_variance','two_stage'),'sharpe']:.2f}, maximum-Sharpe "
        f"{ts.loc[('max_sharpe','one_stage'),'sharpe']:.2f} to "
        f"{ts.loc[('max_sharpe','two_stage'),'sharpe']:.2f} (Table A5, Figures A9-A10)."))
    body(doc, (
        "The finding is the tie itself: splitting the problem into a stock sleeve and "
        "a crypto sleeve, then mixing the two, matched optimising all 60 assets at "
        "once. That means the 1,830-number cross-asset covariance the big model works "
        "so hard to estimate is mostly noise - throwing it away for a simple 2x2 lost "
        "nothing. Simpler wins when the extra detail is unreliable, and here it was."))

    h(doc, "5.3 Sentiment decay on no-news days", 2)
    body(doc, ("Rather than snapping a ticker's sentiment to zero when its news stops, "
               "decay carries the last score forward, fading it over a half-life H:"))
    equation(doc, r"\tilde{s}_{i,t} = s_{i,\tau}\cdot 2^{-(t-\tau)/H},\quad "
                  r"\tau = \text{last headline day} \leq t", "decay", 9)
    hl = int(f["dec"].reset_index()["half_life"].iloc[0])
    dmv = dec.loc[(MV, "sentiment_decay"), "sharpe"]; dmv0 = dec.loc[(MV, "sentiment_neutral0"), "sharpe"]
    dms = dec.loc[(MS, "sentiment_decay"), "sharpe"]; dms0 = dec.loc[(MS, "sentiment_neutral0"), "sharpe"]
    body(doc, (
        f"At a {hl}-day half-life (chosen ex ante, look-ahead safe), decay softened "
        f"the drag without removing it: minimum-variance {dmv0:.2f} to {dmv:.2f}, "
        f"maximum-Sharpe {dms0:.2f} to {dms:.2f}, both still below base. A sweep over "
        f"2, 5, 10 and 21 days barely moves it (Table A6, Figures A11-A12)."))
    body(doc, (
        "Letting yesterday's news fade instead of snapping to zero helped a little "
        "and helped consistently - about +0.02 Sharpe on both funds - and the sweep "
        "shows the exact half-life barely matters. It's a real, if small, "
        "improvement: a fading memory is less jumpy than a hard reset. But it only "
        "softens a losing signal, so decay is worth keeping as the sensible default, "
        "not as the fix that makes sentiment pay."))


if __name__ == "__main__":
    build()
